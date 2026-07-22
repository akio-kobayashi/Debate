from __future__ import annotations

import io
from typing import Any

def _add_distribution(document: Any, question: dict[str, Any]) -> None:
    document.add_heading(question["question"], level=2)
    document.add_paragraph(f"回答数: {question.get('answered', 0)}")
    table = document.add_table(rows=1, cols=3)
    table.style = "Light Shading Accent 1"
    headers = table.rows[0].cells
    headers[0].text = "選択肢"
    headers[1].text = "人数"
    headers[2].text = "割合"
    for item in question.get("distribution", []):
        cells = table.add_row().cells
        cells[0].text = str(item.get("value", ""))
        cells[1].text = str(item.get("count", 0))
        cells[2].text = f"{item.get('percentage', 0)}%"


def _docx_bytes(document: Any) -> bytes:
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_reference_docx(theme: str, reference: dict[str, Any]) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed") from exc
    document = Document()
    document.add_heading("Debate Demo アンケート用参照資料", level=0)
    document.add_paragraph(f"テーマ: {theme}")
    document.add_paragraph(
        "以下はディベート中の発言をLLMが整理した資料です。内容は事実確認済みではありません。"
    )

    motion = reference.get("motion")
    if motion:
        document.add_heading("議題と前提", level=1)
        document.add_paragraph(str(motion))
    for speaker, label in (("A", "A 賛成側"), ("B", "B 反対側")):
        document.add_heading(label, level=1)
        for claim in reference.get("claims", {}).get(speaker, []):
            document.add_heading(
                f"{claim.get('id', speaker)}: {claim.get('title', '論点')}",
                level=2,
            )
            document.add_paragraph(str(claim.get("summary", "")))
            if claim.get("basis"):
                document.add_paragraph(f"根拠・具体例: {claim['basis']}")

    document.add_heading("反論の対応", level=1)
    for rebuttal in reference.get("rebuttals", []):
        document.add_paragraph(
            f"{rebuttal.get('from', '')} → {rebuttal.get('to', '')}: "
            f"{rebuttal.get('summary', '')}"
        )

    document.add_heading("ファシリテーターの整理", level=1)
    summary = reference.get("facilitator_summary", {})
    for key, label in (
        ("agreements", "合意点"),
        ("disagreements", "対立点"),
        ("unresolved", "未解決の点"),
    ):
        values = summary.get(key, [])
        if values:
            document.add_heading(label, level=2)
            for value in values:
                document.add_paragraph(str(value), style="List Bullet")
    return _docx_bytes(document)


def build_analysis_docx(theme: str, analysis: dict[str, Any]) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed") from exc
    document = Document()
    document.add_heading("Debate Demo アンケート分析", level=0)
    document.add_paragraph(f"テーマ: {theme}")
    document.add_paragraph(f"回答者数: {analysis.get('respondent_count', 0)}")

    document.add_heading("設問別集計", level=1)
    for question in analysis.get("questions", []):
        _add_distribution(document, question)

    interpretation = analysis.get("interpretation", "")
    if interpretation:
        document.add_heading("Cによる分析", level=1)
        document.add_paragraph(str(interpretation))
    document.add_paragraph(
        "注: 集計値はサーバーで計算し、Cの文章は集計結果の解釈として生成しています。"
    )
    return _docx_bytes(document)
